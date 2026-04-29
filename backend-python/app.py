from __future__ import annotations

import json
import shutil
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import _Cell, _Row, Table
from docx.text.paragraph import Paragraph
from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename

app = Flask(__name__)
CORS(app)

UPLOAD_DIR = Path(__file__).parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)


@dataclass
class LocatedComment:
    """前端批注在某个段落中的定位结果。"""

    start: int
    end: int
    text: str
    comment: str
    author: str
    initials: str


@app.post("/api/comments/generate")
def generate_comments_docx():
    """接收原始 DOCX 和批注列表，返回带 Word 批注的新 DOCX。"""

    uploaded_file = request.files.get("file")
    if uploaded_file is None:
        return jsonify({"message": "请上传 DOCX 文件"}), 400

    # 前端以 multipart/form-data 传递 comments 字段，内容是 JSON 字符串。
    comments_payload = request.form.get("comments", "[]")
    try:
        comments = json.loads(comments_payload)
    except json.JSONDecodeError:
        return jsonify({"message": "批注数据不是合法 JSON"}), 400

    if not isinstance(comments, list) or len(comments) == 0:
        return jsonify({"message": "至少需要一条批注"}), 400

    # 先用原始文件名判断扩展名，避免中文文件名被 secure_filename 清空后误判。
    original_filename = uploaded_file.filename or "document.docx"
    if not original_filename.lower().endswith(".docx"):
        return jsonify({"message": "仅支持 .docx 文件"}), 400

    # 保存到临时目录时仍使用安全文件名，防止路径穿越或非法字符。
    safe_name = secure_filename(original_filename)
    safe_stem = Path(safe_name).stem or "document"
    filename = f"{safe_stem}.docx"

    # 不使用 TemporaryDirectory 上下文管理器：Windows 下 send_file 读取期间文件会被锁定，
    # 过早清理目录会触发 WinError 32。
    work_dir = Path(tempfile.mkdtemp(dir=UPLOAD_DIR))
    input_path = work_dir / filename
    output_path = work_dir / append_suffix(filename, "-comments")
    uploaded_file.save(input_path)

    document = Document(input_path)
    applied, skipped = apply_comments(document, comments)
    document.save(output_path)

    response = send_file(
        output_path,
        as_attachment=True,
        download_name=output_path.name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    # 等响应真正关闭后再删除临时目录，避免下载流还未读完时清理文件。
    response.call_on_close(lambda: shutil.rmtree(work_dir, ignore_errors=True))
    response.headers["X-Comments-Applied"] = str(applied)
    response.headers["X-Comments-Skipped"] = str(skipped)
    return response


def apply_comments(document: DocumentObject, comments: list[dict]) -> tuple[int, int]:
    """把前端批注定位到段落，并按段落批量写入 Word 批注。"""

    pending = [normalize_comment(comment) for comment in comments]
    paragraph_map: dict[Paragraph, list[LocatedComment]] = {}
    skipped = 0

    # python-docx 的 add_comment 需要锚定到 run；先找到每条批注属于哪个段落。
    paragraphs = list(iter_paragraphs(document))
    for comment in pending:
        location = locate_comment(paragraphs, comment)
        if location is None:
            skipped += 1
            continue

        paragraph, located = location
        paragraph_map.setdefault(paragraph, []).append(located)

    applied = 0
    for paragraph, located_comments in paragraph_map.items():
        applied += rebuild_paragraph_with_comments(document, paragraph, located_comments)

    return applied, len(pending) - applied


def normalize_comment(comment: dict) -> dict:
    """规范化前端提交的批注字段，补齐默认作者信息。"""

    return {
        "selectedText": normalize_text(str(comment.get("selectedText", ""))),
        "occurrence": int(comment.get("occurrence", 0) or 0),
        "comment": str(comment.get("comment", "")).strip(),
        "author": str(comment.get("author", "Reviewer")).strip() or "Reviewer",
        "initials": str(comment.get("initials", "RV")).strip() or "RV",
    }


def locate_comment(
    paragraphs: list[Paragraph],
    comment: dict,
) -> tuple[Paragraph, LocatedComment] | None:
    """按“选中文本 + 第几次出现”在文档段落中定位批注。"""

    selected_text = comment["selectedText"]
    if not selected_text or not comment["comment"]:
        return None

    seen = 0
    for paragraph in paragraphs:
        # 前端 mammoth 预览和 Word 原文的空白可能不同，这里统一折叠空白后匹配。
        # mapping 用来把“折叠空白后的下标”转换回 Word 原段落文本下标。
        paragraph_text, mapping = normalize_text_with_mapping(paragraph.text)
        search_from = 0

        while True:
            start = paragraph_text.find(selected_text, search_from)
            if start == -1:
                break

            if seen == comment["occurrence"]:
                original_start = mapping[start]
                original_end = mapping[start + len(selected_text) - 1] + 1
                return (
                    paragraph,
                    LocatedComment(
                        start=original_start,
                        end=original_end,
                        text=selected_text,
                        comment=comment["comment"],
                        author=comment["author"],
                        initials=comment["initials"],
                    ),
                )

            seen += 1
            search_from = start + len(selected_text)

    return None


def rebuild_paragraph_with_comments(
    document: DocumentObject,
    paragraph: Paragraph,
    located_comments: list[LocatedComment],
) -> int:
    """重建段落 run，使每个被选中文本都有独立 run 可作为批注锚点。"""

    text = paragraph.text
    source_runs = collect_run_spans(paragraph)
    ordered = sorted(located_comments, key=lambda item: item.start)
    non_overlapping: list[LocatedComment] = []
    cursor = 0

    # Word 批注锚点不能互相交叉；重叠批注先跳过，避免生成损坏文档。
    for item in ordered:
        if item.start < cursor:
            continue
        non_overlapping.append(item)
        cursor = item.end

    paragraph.clear()
    position = 0
    runs_for_comments: list[tuple[LocatedComment, list]] = []

    # 将段落拆成“普通文本 run + 批注文本 run”，并复制原 run 的字符格式。
    for item in non_overlapping:
        if item.start > position:
            add_preserved_runs(paragraph, source_runs, position, item.start)

        comment_runs = add_preserved_runs(paragraph, source_runs, item.start, item.end)
        if not comment_runs:
            comment_runs = [paragraph.add_run(text[item.start:item.end])]
        runs_for_comments.append((item, comment_runs))
        position = item.end

    if position < len(text):
        add_preserved_runs(paragraph, source_runs, position, len(text))

    for item, runs in runs_for_comments:
        # python-docx 1.2+ 提供 add_comment，可生成 comments.xml 和文档中的引用锚点。
        document.add_comment(
            runs=runs,
            text=item.comment,
            author=item.author,
            initials=item.initials,
        )

    return len(runs_for_comments)


def iter_paragraphs(parent: DocumentObject | _Cell | _Row | Table) -> Iterable[Paragraph]:
    """递归遍历正文和表格中的段落。"""

    if isinstance(parent, DocumentObject):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            yield from iter_paragraphs(table)
    elif isinstance(parent, Table):
        for row in parent.rows:
            yield from iter_paragraphs(row)
    elif isinstance(parent, _Row):
        for cell in parent.cells:
            yield from iter_paragraphs(cell)
    elif isinstance(parent, _Cell):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            yield from iter_paragraphs(table)


def collect_run_spans(paragraph: Paragraph) -> list[dict]:
    """记录原段落每个 run 在整段文本中的起止位置和格式来源。"""

    spans = []
    position = 0
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue

        start = position
        end = start + len(text)
        spans.append({"start": start, "end": end, "text": text, "run": run})
        position = end

    return spans


def add_preserved_runs(paragraph: Paragraph, source_runs: list[dict], start: int, end: int) -> list:
    """按原 run 边界写入文本片段，并复制对应 run 的字符格式。"""

    new_runs = []
    if start >= end:
        return new_runs

    for span in source_runs:
        overlap_start = max(start, span["start"])
        overlap_end = min(end, span["end"])
        if overlap_start >= overlap_end:
            continue

        relative_start = overlap_start - span["start"]
        relative_end = overlap_end - span["start"]
        new_run = paragraph.add_run(span["text"][relative_start:relative_end])
        copy_run_format(span["run"], new_run)
        new_runs.append(new_run)

    return new_runs


def copy_run_format(source_run, target_run) -> None:
    """复制 run 的底层 rPr，这里包含中文字体、字号、加粗等字符格式。"""

    source_rpr = source_run._r.rPr
    if source_rpr is not None:
        target_run._r.insert(0, deepcopy(source_rpr))


def normalize_text(value: str) -> str:
    """折叠连续空白，降低 HTML 预览文本和 DOCX 原文之间的空白差异。"""

    return " ".join(value.split())


def normalize_text_with_mapping(value: str) -> tuple[str, list[int]]:
    """折叠空白并保留每个规范化字符对应的原始文本下标。"""

    chars = []
    mapping = []
    pending_space_index = None
    emitted_text = False

    for index, char in enumerate(value):
        if char.isspace():
            if emitted_text:
                pending_space_index = index
            continue

        if pending_space_index is not None:
            chars.append(" ")
            mapping.append(pending_space_index)
            pending_space_index = None

        chars.append(char)
        mapping.append(index)
        emitted_text = True

    return "".join(chars), mapping


def append_suffix(filename: str, suffix: str) -> str:
    """给文件名追加后缀，保留原扩展名。"""

    path = Path(filename)
    return f"{path.stem}{suffix}{path.suffix}"


@app.get("/api/health")
def health():
    return jsonify({"ok": True})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
